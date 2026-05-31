from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document as DocxDocument
from pypdf import PdfWriter

from regintel_infrastructure.ingestion import parse_and_chunk


def test_parse_and_chunk_dispatches_docx() -> None:
    document = DocxDocument()
    document.add_paragraph("1. Applicability")
    document.add_paragraph("This Direction applies to all NBFCs registered with the RBI.")
    buffer = BytesIO()
    document.save(buffer)

    chunks = parse_and_chunk("kyc.docx", buffer.getvalue(), uuid4())

    assert len(chunks) > 0
    assert all(chunk.page_number is None for chunk in chunks)


def test_parse_and_chunk_dispatches_pdf() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    chunks = parse_and_chunk("kyc.pdf", buffer.getvalue(), uuid4())

    # A blank page has no extractable text, so no chunks — proves the PDF path ran
    # (not the DOCX path, which would raise trying to parse PDF bytes as a zip).
    assert chunks == []


def test_parse_and_chunk_rejects_unsupported_extensions() -> None:
    with pytest.raises(ValueError, match="unsupported file type"):
        parse_and_chunk("notes.txt", b"irrelevant", uuid4())
