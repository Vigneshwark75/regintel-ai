from io import BytesIO

from docx import Document as DocxDocument
from pypdf import PdfWriter

from regintel_infrastructure.parsing.docx_parser import parse_docx
from regintel_infrastructure.parsing.pdf_parser import parse_pdf


def test_parse_docx_extracts_non_empty_paragraphs() -> None:
    document = DocxDocument()
    document.add_paragraph("1. Applicability")
    document.add_paragraph("")  # blank paragraphs should be dropped
    document.add_paragraph("This Direction applies to all NBFCs.")
    buffer = BytesIO()
    document.save(buffer)

    text = parse_docx(buffer.getvalue())

    assert "1. Applicability" in text
    assert "This Direction applies to all NBFCs." in text


def test_parse_pdf_returns_one_entry_per_page() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)

    pages = parse_pdf(buffer.getvalue())

    assert [p.page_number for p in pages] == [1, 2]
