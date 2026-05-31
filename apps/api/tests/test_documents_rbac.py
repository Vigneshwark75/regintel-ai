from io import BytesIO

from api.dependencies import get_ingest_document_use_case
from api.main import app
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

client = TestClient(app)


class FakeIngestUseCase:
    def __init__(self) -> None:
        self.calls: list[tuple[object, list[object]]] = []

    async def execute(self, document: object, chunks: list[object]) -> None:
        self.calls.append((document, chunks))


def _get_token(username: str, password: str) -> str:
    response = client.post("/auth/token", data={"username": username, "password": password})
    return str(response.json()["access_token"])


def _make_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("1. Applicability")
    document.add_paragraph("This Direction applies to all NBFCs registered with the RBI.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_UPLOAD_FORM = {
    "title": "Master Direction on KYC",
    "document_type": "master_direction",
    "reference_number": "RBI/DBR/2016-17/18",
    "issued_date": "2016-02-25",
}


def test_upload_is_forbidden_for_a_role_that_does_not_manage_documents() -> None:
    token = _get_token("risk", "risk-demo-password")

    response = client.post(
        "/documents",
        headers={"Authorization": f"Bearer {token}"},
        files={"file": ("kyc.docx", _make_docx_bytes())},
        data=_UPLOAD_FORM,
    )

    assert response.status_code == 403


def test_upload_succeeds_for_a_compliance_officer() -> None:
    fake_use_case = FakeIngestUseCase()
    app.dependency_overrides[get_ingest_document_use_case] = lambda: fake_use_case
    try:
        token = _get_token("compliance", "compliance-demo-password")

        response = client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("kyc.docx", _make_docx_bytes())},
            data=_UPLOAD_FORM,
        )

        assert response.status_code == 201
        assert response.json()["chunk_count"] > 0
        assert len(fake_use_case.calls) == 1
    finally:
        app.dependency_overrides.pop(get_ingest_document_use_case, None)


def test_upload_succeeds_for_a_cro() -> None:
    fake_use_case = FakeIngestUseCase()
    app.dependency_overrides[get_ingest_document_use_case] = lambda: fake_use_case
    try:
        token = _get_token("cro", "cro-demo-password")

        response = client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("kyc.docx", _make_docx_bytes())},
            data=_UPLOAD_FORM,
        )

        assert response.status_code == 201
    finally:
        app.dependency_overrides.pop(get_ingest_document_use_case, None)


def test_upload_rejects_an_unsupported_file_type() -> None:
    fake_use_case = FakeIngestUseCase()
    app.dependency_overrides[get_ingest_document_use_case] = lambda: fake_use_case
    try:
        token = _get_token("compliance", "compliance-demo-password")

        response = client.post(
            "/documents",
            headers={"Authorization": f"Bearer {token}"},
            files={"file": ("notes.txt", b"plain text")},
            data=_UPLOAD_FORM,
        )

        assert response.status_code == 400
        assert fake_use_case.calls == []
    finally:
        app.dependency_overrides.pop(get_ingest_document_use_case, None)
