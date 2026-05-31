from io import BytesIO

import pytest
from api.main import app
from docx import Document as DocxDocument
from fastapi.testclient import TestClient

pytestmark = pytest.mark.integration


def _make_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_paragraph("1. Customer Due Diligence")
    document.add_paragraph(
        "Regulated entities shall retain customer due diligence records for a "
        "minimum of five years from the date of the transaction or the end of "
        "the relationship, whichever is later."
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_full_flow_upload_then_ask_returns_a_grounded_answer() -> None:
    with TestClient(app) as client:
        token_response = client.post(
            "/auth/token", data={"username": "compliance", "password": "compliance-demo-password"}
        )
        assert token_response.status_code == 200
        headers = {"Authorization": f"Bearer {token_response.json()['access_token']}"}

        upload_response = client.post(
            "/documents",
            headers=headers,
            files={"file": ("kyc.docx", _make_docx_bytes())},
            data={
                "title": "Master Direction on KYC",
                "document_type": "master_direction",
                "reference_number": "RBI/DBR/2016-17/18",
                "issued_date": "2016-02-25",
            },
        )
        assert upload_response.status_code == 201
        document_id = upload_response.json()["document_id"]

        ask_response = client.post(
            "/ask",
            headers=headers,
            json={"question": "How long must customer due diligence records be retained?"},
        )

        assert ask_response.status_code == 200
        body = ask_response.json()
        assert "five" in body["answer"].lower() or "5" in body["answer"]
        assert len(body["citations"]) > 0
        assert any(c["document_id"] == document_id for c in body["citations"])
